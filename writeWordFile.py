from docx import Document
from docx.shared import RGBColor,Pt,Cm
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from wordTopdf import convert_to_pdf
import os


class BewerbungDocx:
	def __init__(self):
		#self.document = Document()
		self.now = datetime.now()

		self.center = WD_ALIGN_PARAGRAPH.CENTER
		self.right = WD_ALIGN_PARAGRAPH.RIGHT
		self.left = WD_ALIGN_PARAGRAPH.LEFT
		self.justify = WD_ALIGN_PARAGRAPH.JUSTIFY

	def createPage(self):
		self.document = Document()
		return self.document

	def save_file(self,name: str):

		try:
			file = os.path.isfile(f"{name}.docx")
			if not file:
				self.document.save(f'{name}.docx')

			else:
				print("the file already created")

		except Exception:
			print("maybe Your File is open close it pls")

	def write_paragraph(self,align,letters: str = "",size: int = 9,r: int = 0,g: int = 0,b: int = 0,font_family: str = "Arial"):
		paragraph = self.document.add_paragraph()
		paragraph.alignment = align
		section = paragraph.add_run(f"{letters}")
		section.font.size = Pt(size)
		section.font.color.rgb = RGBColor(r,g,b)
		section.font.name = f"{font_family}"

	def write_heading(self,align,letters: str = "",size: int = 9,r: int = 0,g: int = 0,b: int = 0,font_family: str = "Arial",l=0):
		paragraph = self.document.add_heading(level=l)
		paragraph.alignment = align
		section = paragraph.add_run(f"{letters}")
		section.font.size = Pt(size)
		section.font.color.rgb = RGBColor(r,g,b)
		section.font.name = f"{font_family}"

	def page_margin(self,top=0.79,bottom=0.79,left=1,right=1):
		sections = self.document.sections
		for section in sections:
			section.top_margin = Cm(top)
			section.bottom_margin = Cm(bottom)
			section.left_margin = Cm(left)
			section.right_margin = Cm(right)

	def convertToPdf(self,filename):
		convert_to_pdf(filename)


# x = BewerbungDocx()
# x.convertToPdf()
